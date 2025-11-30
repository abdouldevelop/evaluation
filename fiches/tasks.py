
from celery import shared_task
from django.conf import settings
from docx import Document
from docx2pdf import convert
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import pythoncom
import os
import tempfile
from fiches.models import Evaluation

@shared_task
def generate_evaluation_pdf(evaluation_id):
      evaluation = Evaluation.objects.get(id=evaluation_id)
      agent = evaluation.agent

      doc_path = os.path.join(settings.MEDIA_ROOT, 'EVALU.docx')
      if not os.path.exists(doc_path):
          return {"status": "error", "message": "Fichier modèle Word introuvable"}

      doc = Document(doc_path)
      total_semestre = float(evaluation.moyenne_rendement) * 2 + float(evaluation.moyenne_comportement)
      moyenne_semestre = total_semestre / 3

      remplacements = {
          "agent_nom": f"{agent.nom} {agent.prenoms}",
          "agent_matricule": agent.matricule,
          "agent_date_embauche": agent.date_embauche.strftime("%d/%m/%Y"),
          "agent_categorie": agent.categorie,
          "agent_direction": agent.direction.nom,
          "agent_sous_direction": agent.sous_direction if agent.sous_direction else "N/A",
          "agent_service": agent.service if agent.service else "N/A",
          "agent_poste": agent.poste if hasattr(agent, 'poste') else "N/A",
          "agent_tenu_le": agent.tenu_depuis.strftime("%d/%m/%Y") if hasattr(agent, 'tenu_depuis') else "N/A",
          "evaluations.0.connaissances": str(evaluation.connaissances),
          "evaluations.0.initiative": str(evaluation.initiative),
          "evaluations.0.rendement": str(evaluation.rendement),
          "evaluations.0.respect_objectifs": str(evaluation.respect_objectifs),
          "evaluations.0.total_rendement": str(evaluation.connaissances + evaluation.initiative + evaluation.rendement + evaluation.respect_objectifs),
          "evaluations.0.moyenne_rendement": f"{evaluation.moyenne_rendement:.3f}",
          "evaluations.0.civisme": str(evaluation.civisme),
          "evaluations.0.service_public": str(evaluation.service_public),
          "evaluations.0.relations_humaines": str(evaluation.relations_humaines),
          "evaluations.0.discipline": str(evaluation.discipline),
          "evaluations.0.ponctualité": str(evaluation.ponctualité),
          "evaluations.0.assiduite": str(evaluation.assiduite),
          "evaluations.0.tenue": str(evaluation.tenue),
          "evaluations.0.total_comportement": str(evaluation.civisme + evaluation.service_public + evaluation.relations_humaines + evaluation.discipline + evaluation.ponctualité + evaluation.assiduite + evaluation.tenue),
          "evaluations.0.moyenne_comportement": f"{evaluation.moyenne_comportement:.2f}",
          "E.MR": str(float(evaluation.moyenne_rendement) * 2),
          "E.MC": str(float(evaluation.moyenne_comportement)),
          "E.TOU": str(total_semestre),
          "M.0.T": str(moyenne_semestre),
          "evaluations.0.avis_directeur": evaluation.avis_directeur if evaluation.avis_directeur else "Aucun avis",
          "evaluations.0.avis_agent": evaluation.avis_agent if evaluation.avis_agent else "Aucun avis",
      }

      mots_cles_a_styliser = ["moyenne", "total", "e.mr", "e.mc", "e.tou", "m.0.t"]

      for para in doc.paragraphs:
          for key, value in remplacements.items():
              if key in para.text:
                  para.text = para.text.replace(key, str(value))
                  if any(mot_cle in key.lower() for mot_cle in mots_cles_a_styliser):
                      para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                      for run in para.runs:
                          run.font.size = Pt(12)

      for table in doc.tables:
          for row in table.rows:
              for cell in row.cells:
                  for key, value in remplacements.items():
                      if key in cell.text:
                          cell.text = cell.text.replace(key, str(value))
                          if any(mot_cle in key.lower() for mot_cle in mots_cles_a_styliser):
                              for paragraph in cell.paragraphs:
                                  paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                  for run in paragraph.runs:
                                      run.bold = True
                                      run.font.size = Pt(12)
                          elif "evaluation" in key.lower():
                              for paragraph in cell.paragraphs:
                                  paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                  for run in paragraph.runs:
                                      run.font.size = Pt(12)

      with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_word_file:
          temp_word_path = temp_word_file.name
          doc.save(temp_word_path)

      try:
          pythoncom.CoInitialize()
          with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_pdf_file:
              temp_pdf_path = temp_pdf_file.name
          convert(temp_word_path, temp_pdf_path)

          if not os.path.exists(temp_pdf_path):
              return {"status": "error", "message": "Conversion en PDF échouée"}

          with open(temp_pdf_path, "rb") as pdf_file:
              evaluation.fiche_pdf.save(
                  f"Fiche_Evaluation_{agent.matricule}_{evaluation.annee}_{evaluation.semestre}.pdf",
                  pdf_file,
                  save=True
              )

          return {"status": "success", "message": "PDF généré et stocké"}

      except Exception as e:
          return {"status": "error", "message": str(e)}

      finally:
          pythoncom.CoUninitialize()
          if os.path.exists(temp_word_path):
              os.remove(temp_word_path)
          if os.path.exists(temp_pdf_path):
              os.remove(temp_pdf_path)